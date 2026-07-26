#!/usr/bin/env python3
"""Build the Chapter 12 spine-seed decision packet as DOCX.

Companion to pipeline/ch12-spine-seeds.md, which remains the diffable record.
Requirements from pipeline/spine-seed-protocol.md -> Delivery format:
  * real Word heading styles so the navigation pane works
  * each decision in a shaded, numbered box with an adjacent empty answer field
  * an index of all decisions at the front
  * cuts and closed items in visually distinct callouts
  * no inline code formatting or markdown syntax in the rendered text
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "line-edit-packets" / "Chapter_12_Spine_Seeds.docx"

INK = RGBColor(0x11, 0x11, 0x11)
NAVY = RGBColor(0x0B, 0x25, 0x45)
GREY = RGBColor(0x55, 0x55, 0x55)
DECISION_FILL = "FFF3D6"   # warm — action required
ANSWER_FILL = "FFFFFF"
CLOSED_FILL = "E8EEF4"     # cool — recorded, no action
FINDING_FILL = "EDF4EC"    # green — finding
SEED_FILL = "F7F7F5"


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hex_fill)
    tcPr.append(el)


def borders(table, sz=6, color="BBBBBB"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    el = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color)
        el.append(e)
    tblPr.append(el)


def fix_widths(table, widths):
    """Force a fixed layout with explicit column widths.

    python-docx alone does not hold column widths: Word and LibreOffice both
    re-autofit unless the layout is declared fixed, the grid carries the widths,
    and every cell repeats them.
    """
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        table._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for wd in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(wd * 1440)))
        grid.append(gc)
    table._tbl.insert(1, grid)

    for row in table.rows:
        for i, wd in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(wd)


def no_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def keep_with_next(cell):
    for p in cell.paragraphs:
        p.paragraph_format.keep_with_next = True


def run(p, text, *, bold=False, italic=False, size=10.5, color=INK):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return r


def para(doc, text="", *, bold=False, italic=False, size=10.5, color=INK,
         before=0, after=6, indent=None, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent is not None:
        pf.left_indent = Inches(indent)
    if align is not None:
        p.alignment = align
    if text:
        run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def rich(doc, segments, *, size=10.5, before=0, after=6, indent=None):
    """segments: list of (text, style) where style in {'', 'b', 'i', 'bi'}."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent is not None:
        pf.left_indent = Inches(indent)
    for text, style in segments:
        run(p, text, bold="b" in style, italic="i" in style, size=size)
    return p


def bullet(doc, segments, *, size=10.5, indent=0.25):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.left_indent = Inches(indent + 0.25)
    for text, style in segments:
        run(p, text, bold="b" in style, italic="i" in style, size=size)
    return p


def heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = NAVY
        r.font.name = "Calibri"
        if level == 1:
            r.font.size = Pt(16)
        elif level == 2:
            r.font.size = Pt(13)
        else:
            r.font.size = Pt(11.5)
    h.paragraph_format.space_before = Pt(14 if level < 3 else 10)
    h.paragraph_format.space_after = Pt(4)
    return h


def callout(doc, label, body_lines, fill, *, label_color=NAVY):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(t)
    fix_widths(t, [6.5])
    c = t.cell(0, 0)
    shade(c, fill)
    first = c.paragraphs[0]
    first.paragraph_format.space_after = Pt(3)
    run(first, label, bold=True, size=10, color=label_color)
    for line in body_lines:
        p = c.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        if isinstance(line, str):
            run(p, line, size=10)
        else:
            for text, style in line:
                run(p, text, bold="b" in style, italic="i" in style, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def decision(doc, number, title, body_lines, *, answer_lines=3):
    t = doc.add_table(rows=2, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(t, sz=8, color="D9A93C")
    fix_widths(t, [6.5])
    for r in t.rows:
        no_split(r)
    top = t.cell(0, 0)
    shade(top, DECISION_FILL)
    h = top.paragraphs[0]
    h.paragraph_format.space_after = Pt(4)
    run(h, f"DECISION {number} — {title}", bold=True, size=10.5, color=NAVY)
    for line in body_lines:
        p = top.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        if isinstance(line, str):
            run(p, line, size=10)
        else:
            for text, style in line:
                run(p, text, bold="b" in style, italic="i" in style, size=10)
    keep_with_next(top)
    bot = t.cell(1, 0)
    shade(bot, ANSWER_FILL)
    lab = bot.paragraphs[0]
    lab.paragraph_format.space_after = Pt(2)
    run(lab, "Your answer:", bold=True, size=9, color=GREY)
    for _ in range(answer_lines):
        p = bot.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        run(p, "", size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def datatable(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(t)
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    for i, htxt in enumerate(headers):
        c = t.cell(0, i)
        shade(c, "E4E9EF")
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run(p, htxt, bold=True, size=9.5, color=NAVY)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            emphasise = str(val).startswith("*")
            txt = str(val)[1:] if emphasise else str(val)
            run(p, txt, bold=emphasise, size=9.5)
    fix_widths(t, widths)
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def seed(doc, fields):
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(t, color="CCCCCC")
    for label, segments in fields:
        cells = t.add_row().cells
        shade(cells[0], SEED_FILL)
        lp = cells[0].paragraphs[0]
        lp.paragraph_format.space_after = Pt(2)
        run(lp, label, bold=True, size=9.5, color=NAVY)
        vp = cells[1].paragraphs[0]
        vp.paragraph_format.space_after = Pt(2)
        for text, style in segments:
            run(vp, text, bold="b" in style, italic="i" in style, size=10)
        no_split(t.rows[-1])
    fix_widths(t, [1.3, 5.2])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def build():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(0.9)

    # footer page numbers
    for s in doc.sections:
        fp = s.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(fp, "Chapter 12 Spine Seeds — page ", size=8.5, color=GREY)
        r = fp.add_run()
        r.font.size = Pt(8.5)
        r.font.color.rgb = GREY
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for el in (begin, instr, end):
            r._r.append(el)

    # ---------------- title ----------------
    tp = doc.add_paragraph()
    tp.paragraph_format.space_after = Pt(2)
    run(tp, "Chapter 12 — Emotion, Stress & Coping", bold=True, size=22, color=NAVY)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(10)
    run(sp, "Spine Seeds — for instructor edit", size=13, color=GREY)

    para(doc, "Drafted 2026-07-26. Nine decisions. Nothing here is prose for the book.",
         size=10, color=GREY, after=8)

    callout(doc, "How to use this", [
        "Edit in place. Cross out what is wrong, rewrite the payoff lines in your own "
        "words, and mark each decision keep, cut, or change. When it comes back I draft "
        "prose against it.",
        [("The drafted payoff lines are not proposals to approve. They exist so you have "
          "something to react to instead of authoring cold. Rewriting one is the expected "
          "outcome, not a correction.", "i")],
        "Single-packet format at your direction. Everything I could decide myself is in "
        "Recorded, Not Asked near the end rather than in a box.",
    ], CLOSED_FILL)

    # ---------------- index ----------------
    heading(doc, "The nine decisions", 1)
    para(doc, "Scope before you start. Page order follows the chapter, not this list.",
         size=10, color=GREY, after=6)
    datatable(doc,
              ["#", "Decision", "Why it is yours"],
              [["1", "Spine question wording", "Register and student framing"],
               ["2", "Chapter engine sentence", "Voice"],
               ["3", "Social buffering thread — install it?", "Adding a thread"],
               ["4", "Compression thread — install it?", "Adding a thread"],
               ["5", "Action-readiness — name it or leave it implicit?", "Adding a thread"],
               ["6", "The Patient S.M. payoff line", "Sits on a high-severity must-correct"],
               ["7", "Seven prunes, 184 words", "Every prune is yours"],
               ["8", "Cognitive light cone must-preserve conflict", "Must-preserve, plus a spine amendment"],
               ["9", "Net +136 words against a protocol rule", "Scope"]],
              widths=[0.4, 3.5, 2.6])

    # ---------------- diagnostic ----------------
    heading(doc, "Why this chapter needs seeds — and why the reason is not Chapter 11's", 1)
    para(doc, "Chapter 11 was a construct inventory: four Layer 3 threads scored zero in a "
              "chapter the spine assigns them to. Chapter 12 is not that. Its threads are "
              "present and mostly well taught.")

    heading(doc, "Thread-count diagnostic", 2)
    datatable(doc, ["Layer 3 thread", "Count", "Status"],
              [["allostatic load / allostasis", "13 / 13", "present, well taught"],
               ["appraisal", "18", "present, well taught"],
               ["core affect", "13", "present, well taught"],
               ["amygdala / Patient S.M.", "13 / 6", "present; CO2 boundary installed"],
               ["Barrett / constructionism", "6", "present; named once as chosen lens"],
               ["cognitive light cone", "5", "present; now attributed to Levin (2019)"],
               ["hedonic adaptation", "4", "present; descriptive only, correctly"],
               ["*action-readiness / action tendency", "*1 / 2", "*table cell and review answer only"],
               ["*mattering", "*1", "*appears only inside a disclaimer"],
               ["*compression / lossy", "*0", "*spine's central lens; section 7 wants it everywhere"],
               ["*belonging", "*0", "*the Chapter 11 thread this chapter reconnects"],
               ["*calibration", "*0", ""],
               ["*social baseline / Coan", "*0", "*expected; promoted 2026-07-25"],
               ["dopamine", "0", "correctly absent — see Finding 3"]],
              widths=[2.9, 0.9, 2.7])

    heading(doc, "The finding: the payoff stopped arriving", 2)
    rich(doc, [("The repair pass ran the day after the audit and cut the chapter from 11,050 "
                "words to 6,120. Body prose fell from about 6,198 to ", ""),
               ("3,653", "b"),
               (" — the low end of the 3,500 to 5,000 target, not over it. What survived is "
                "disproportionately boundary prose.", "")])
    rich(doc, [("Forty per cent of body paragraphs now end on a negation, against 25 per cent "
                "for Chapter 11 and 22 per cent for Chapter 8. ", ""),
               ("That proxy overstates the problem by about half and should not be quoted as "
                "it stands.", "b"),
               (" Hand-checking, six of eleven flagged paragraphs are correctly calibrated, "
                "because in this chapter a negation is often the teaching content. “Core "
                "affect is not a specific emotion” is the concept, not a hedge on it. "
                "“It is not a single fear button” arrives after a positive claim and "
                "is a good line. Five are genuine failures, and they fail identically: the "
                "paragraph never states its conclusion.", "")])
    datatable(doc, ["Line", "Paragraph", "How it fails"],
              [["L55", "body budget",
                "Four sentences, three about what the metaphor is not. A student learns the "
                "tank is not real, four times, and never learns what the metaphor is for."],
               ["L76", "sleep, hunger, isolation",
                "Hedged before the claim lands, then a study, then closes on not turning "
                "fatigue into a diagnosis."],
               ["L88", "constructionism boundary",
                "Entirely meta — about the chapter's epistemics, not about emotion — and ends "
                "by announcing it will not repeat itself."],
               ["L121", "journaling, friendship, therapy",
                "Two sentences, both negative. Never says what those things do."],
               ["L177", "oxytocin and Heinrichs",
                "Ends by dismantling its own evidence. Section 4 then says social support "
                "helps, with no mechanism."]],
              widths=[0.55, 1.6, 4.35])
    para(doc, "Two of these violate the AGENTS style rule against epistemological "
              "throat-clearing close to verbatim, ending on what a study does not establish.")

    heading(doc, "Root cause — it changes the fix", 2)
    rich(doc, [("The audit's own revision constraints say: ", ""),
               ("do not remove body budget or the cognitive light cone without replacing "
                "their teaching functions.", "i"),
               (" The repair pass correctly removed three overclaims — the body-budget "
                "mechanism claim, social support predicting mortality, and the post-finals "
                "explanation — and replaced none of their teaching functions.", "")])
    rich(doc, [("So the fix is drafting payoffs, not deleting caveats.", "b"),
               (" That is why this packet has no register section: the work rides inside the "
                "payoff-line field the protocol already specifies. Recorded as a general rule "
                "in AGENTS, under Audits are diagnostic, not authoritative.", "")])

    # ---------------- findings ----------------
    heading(doc, "Findings that are not seeds", 1)

    callout(doc, "FINDING 1 — the social baseline anchor was the wrong study, and the search found better", [
        "The spine names Coan, Schaefer and Davidson (2006) as the anchor. It is 16 married "
        "women, almost all white, highly satisfied, one city. Introducing it would land three "
        "paragraphs after this chapter tells students a 37-person oxytocin experiment does not "
        "establish a general mechanism — the Chapter 11 Asch inconsistency in reverse.",
        [("Selected instead: Hostinar, Sullivan and Gunnar (2014), Psychological Bulletin "
          "140(1), 256–282.", "b"),
         (" A translational review of social buffering of the HPA axis across animal models "
          "and human studies, infancy to adulthood. Review-level rather than one small study; "
          "HPA, which Section 3 already teaches and Chapter 3 introduced; developmental, so it "
          "is the bridge this chapter already promises; and it names oxytocin systems as a "
          "putative mediator, which turns the orphaned Heinrichs paragraph from a debunk into "
          "a small study bounded by a review.", "")],
        "Available if you would rather have an effect size than a review: Thorsteinsson and "
        "James (1999), 22 studies, 56 effect sizes, N=1,167, reduced reactivity across heart "
        "rate, both blood pressures, and cortisol. Strongest causal warrant here; no mechanism "
        "story, so it makes the claim true without making it memorable.",
        "Rejected: Coan et al. (2017) fMRI, n=110 — best findings for this audience, but adds "
        "an imaging paradigm to a chapter already carrying a lesion case and two fMRI claims, "
        "in the one section the audit flagged for thin citation. Beckes and Sbarra (2022) — the "
        "authors' own framing is that evidence is mounting, which tells you the payoff is thin. "
        "Hostinar, Johnson and Gunnar (2015) — parental buffering losing potency at mid-puberty "
        "is striking for readers who just came through it, but N=40 and 41 and the adolescent "
        "result is a null at that size.",
    ], FINDING_FILL)

    callout(doc, "FINDING 2 — the chapter promises this material twice and cannot keep the promise", [
        "Where This Fits says Chapter 10 showed how relationships can support regulation "
        "across development. The Connections table says the same. Chapter 10 does deliver a "
        "version, in your voice:",
        [("For a while, the caregiver's regulated body is the infant's regulation system, "
          "borrowed from the outside. Over months and years, the child internalizes capacities "
          "that were first shared.", "i")],
        "That sentence is social baseline theory, for infants. The gap is not a missing thread. "
        "It is a missing continuation — which makes this a repair, not an addition.",
    ], FINDING_FILL)

    callout(doc, "FINDING 3 — the spine's own Chapter 12 note overclaims relative to its own drafting rules", [
        "The hedonic-treadmill note prescribes the dopamine anticipation mechanism. The spine's "
        "section 7 forbids generalizing dopamine reward-prediction error to other emotional "
        "systems, and the audit lists that exact mechanism under must-correct. The chapter chose "
        "section 7 and is right. The spine note needs amending, not the chapter. Not routed as a "
        "decision because it changes nothing here — flagged so it is not re-proposed next pass.",
    ], CLOSED_FILL)

    callout(doc, "FINDING 4 — Grupe and Nitschke verifies, and should still be declined", [
        "The spine flags it verify before citing. It verifies: Grupe and Nitschke (2013), "
        "Nature Reviews Neuroscience 14, 488–501, real and well cited. But it is about "
        "pathological anxiety, so using it here imports the diagnostic frame the audit's "
        "Chapter 13 bridge explicitly bars. Recommend closing the flag by declining it in "
        "Chapter 12 and noting it as a Chapter 13 candidate. A check that cleared the citation "
        "and failed the use.",
    ], CLOSED_FILL)

    callout(doc, "FINDING 5 — Levin (2019) is real, which makes two of our documents stale", [
        "Levin defines cognitive light cone as the outer spatial and temporal boundary of the "
        "largest goal a system can work toward. So the audit's requirement that the term keep "
        "its course-specific label, and the spine's description of it as a course metaphor and "
        "not a standard research term, are both overtaken by a correct attribution. See "
        "Decision 8.",
    ], FINDING_FILL)

    # ---------------- chapter level ----------------
    heading(doc, "Chapter level", 1)

    heading(doc, "The spine question", 2)
    rich(doc, [("From the spine: ", ""), ("How does the body decide what matters?", "bi")])
    para(doc, "Unlike Chapter 11's, this one has no abstraction problem — matters is student "
              "language and it is the chapter's own thread word. Its weakness is coverage: it "
              "frames Sections 1 and 2 cleanly and leaves stress and coping to follow on. Two "
              "alternatives keep the claim and add the second clause that earned its keep in "
              "Chapter 11:")
    bullet(doc, [("How does your body decide what matters — and why won't it always stand down?", "i")])
    bullet(doc, [("Why does your body react to things that haven't happened yet?", "i")])
    decision(doc, 1, "Spine question wording", [
        "Keep the spine's wording, take one of the alternatives, or write your own.",
        "It lands at the end of the smoke-alarm opener. The audit protects that opener, so the "
        "question sharpens its last paragraph rather than replacing it.",
    ])

    heading(doc, "The one-sentence chapter engine", 2)
    para(doc, "Every section should be traceable to this. If a passage cannot be, it is catalog.")
    rich(doc, [("Nothing about you gets to wait for certainty. Your body prepares for what "
                "matters before you have decided what matters — and what you feel is that "
                "preparation.", "bi")], indent=0.3)
    para(doc, "It maps all four sections: anticipatory regulation, the same preparation "
              "categorized into a specific emotion, preparation that never gets called off, "
              "and deciding what to do with it. It also answers the spine question by "
              "reframing it — the body decides before you do.")
    decision(doc, 2, "Chapter engine sentence", ["Does this sound like you? Rewrite it if not."])

    # ---------------- section seeds ----------------
    heading(doc, "Section seeds", 1)

    heading(doc, "Opener — “Emotions happen to you”", 2)
    seed(doc, [
        ("Spine claim", [("The reflex model is right about the speed and wrong about the "
                          "history. Something decides the alarm is worth sounding, and that "
                          "something has been learning for years.", "")]),
        ("Lead with", [("The smoke alarm, unchanged. The audit protects the image, the direct "
                        "address, and the constructed-does-not-mean-chosen distinction. It is "
                        "already the best cold open in the chapter.", "")]),
        ("Arc", [("Event fires a circuit, out comes a feeling → same event, two people, "
                  "two feelings → it is not reacting to the world but to what the world "
                  "might mean.", "")]),
        ("Payoff line", [("The alarm in your kitchen reacts to smoke. You react to what smoke "
                          "might mean — and nobody arrives at that with the same history.", "i")]),
        ("Prune", [("None. This section is 283 words and earns them.", "")]),
        ("Watch", [("Must-preserve. Do not restore “emotions are constructed predictions, "
                    "not hardwired triggers” — the spine's section 7 bars presenting "
                    "construction as settled, and the removed version asserted it.", "")]),
    ])

    heading(doc, "Section 1 — The Regulated Body", 2)
    seed(doc, [
        ("Spine claim", [("Regulation runs ahead of need. Allostasis is prediction pointed "
                          "inward — the machinery Chapter 4 aimed at perception and Chapter 8 "
                          "at memory, aimed at the body. This is where the compression lens "
                          "installs.", "")]),
        ("Lead with", [("The thermostat that learns your schedule, with its verbs back "
                        "(Reclaim 2). It currently arrives after two paragraphs of "
                        "definition.", "")]),
        ("Arc", [("You assume the body corrects errors → it moves first, before the error "
                  "→ feeling is part of getting ready, not a report on what happened.", "")]),
        ("Payoff line", [("By the time you feel anything, your body has already moved. Feeling "
                          "is not the report on what happened — it is part of getting ready "
                          "for what might.", "i")]),
        ("Compression seam", [("Two numbers — pleasant or not, wound up or not — standing in "
                               "for everything happening in a body. That is a summary, and "
                               "summaries leave things out. It is also why you can be certain "
                               "you feel terrible and have no idea why.", "i"),
                              ("  (Decision 4)", "b")]),
        ("Prune", [("P2 — the two interchangeability examples at L55, 22 words. P7 — L69's "
                    "closing sentence, 17 words; it restates the sentence before it.", "")]),
        ("Watch", [("Body budget is a metaphor and nothing else. The must-preserve ledger, the "
                    "figure 12.1 decision, and the must-not-happen list all bar one measurable "
                    "currency, one literal tank, and quasi-diagnostic use. The compression seam "
                    "must describe core affect as a summary, never as a meter. Reclaim 2 "
                    "touches a must-preserve row.", "")]),
    ])

    heading(doc, "Section 2 — How Does a Feeling Become an Emotion?", 2)
    seed(doc, [
        ("Spine claim", [("The same bodily state becomes different emotions depending on which "
                          "concept the brain brings to it — and the concept changes what you do "
                          "next. The spine's formulation: concepts predict meaning, emotions "
                          "predict mattering.", "")]),
        ("Lead with", [("The granularity examples. Anxious, frustrated and disappointed "
                        "currently arrive at the end of the section, after the theory table and "
                        "Patient S.M. They are the concrete hook and they are in the wrong "
                        "position.", "")]),
        ("Arc", [("You have a feeling and a word for it → the word changes the action "
                  "→ precision is not vocabulary, it is a different instruction.", "")]),
        ("Payoff line", [("“I feel bad” tells you nothing about what to do next. "
                          "“I'm disappointed” tells you to revise an expectation. "
                          "“I'm frustrated” tells you to change the approach. Same "
                          "discomfort, different instructions.", "i")]),
        ("Action-readiness", [("The payoff line already carries it — instructions is the "
                               "action-readiness claim in student language. One further "
                               "sentence if you want it named: ", ""),
                              ("Every emotion arrives with something already leaning: fear "
                               "toward getting out, anger toward pushing back, shame toward "
                               "getting small.", "i"),
                              ("  (Decision 5)", "b")]),
        ("Prune", [("P1 — L88 entire, 59 words. The largest single prune and the one with a "
                    "consequence; see Decision 7.", "")]),
        ("Watch", [("Section 7: do not say emotion and cognition are separate systems; do not "
                    "say concepts are explicit and emotions implicit; do not present "
                    "construction as settled. Constructionism must be named once as the "
                    "chapter's chosen lens — if P1 goes, that has to land somewhere else.", "")]),
    ])

    heading(doc, "Section 2, Classic Study — Patient S.M.", 2)
    seed(doc, [
        ("Spine claim", [("The case separates knowing from minding. She can state that a "
                          "situation should be frightening without the state arriving.", "")]),
        ("Lead with", [("The snakes. The walkthrough currently opens on a question about "
                        "bilateral damage.", "")]),
        ("Arc", [("She knows the facts → nothing follows from them → and then CO2 "
                  "produced fear anyway, so there is more than one route in.", "")]),
        ("Payoff line", [("She can tell you the snake is dangerous. She just doesn't mind. "
                          "Knowing and minding turn out to be two different jobs — and the CO2 "
                          "result shows the amygdala is not the only way into the second one.", "i")]),
        ("Watch", [("Highest-risk line in the packet. See Decision 6.", "")]),
    ])
    decision(doc, 6, "The Patient S.M. payoff line", [
        [("The audit required removing ", ""),
         ("“Take it out, and the facts survive; the mattering does not,”", "i"),
         (" and lists Patient S.M. as evidence that amygdala removal eliminates felt mattering "
          "or all fear under must-correct.", "")],
        "My line attributes the knowing-and-minding distinction to the pattern of findings "
        "rather than making the amygdala the converter, and the CO2 clause does the bounding "
        "work in the same breath. I believe it is compliant. It is close enough to a "
        "high-severity must-correct that I am not deciding it myself.",
        "Keep as drafted, rewrite, or drop the distinction and leave the walkthrough factual.",
    ])

    heading(doc, "Section 3 — Stress", 2)
    seed(doc, [
        ("Spine claim", [("Stress is mobilization for a predicted demand. What makes human "
                          "stress different is not the machinery but the range of things that "
                          "can count as a demand.", "")]),
        ("Lead with", [("The zebra (Reclaim 3). The section currently opens on the Lazarus and "
                        "Folkman definition. Sapolsky's zebra contrast is a voice-brief "
                        "signature example and it now survives only in Further Reading.", "")]),
        ("Arc", [("The zebra mobilizes fully, then stands down → yours doesn't, because "
                  "your threats don't resolve → payoff.", "")]),
        ("Payoff line", [("The zebra's problem ends. Yours is next Tuesday — and it will still "
                          "be next Tuesday tomorrow.", "i")]),
        ("Prune", [("P6 — L137's closing scope clause, 20 words. The goal-horizon ladder in the "
                    "same paragraph can also lose a sentence; recorded, not asked.", "")]),
        ("Watch", [("Overlapping, never identical. GAS stays a historical heuristic. Cortisol "
                    "does not cause the immediate heart-rate rise. Reclaim 3 needs one word "
                    "changed from the removed version — the same ancient mobilization machinery "
                    "becomes overlapping — or it reintroduces a must-correct. Light cone "
                    "attribution: Levin, and see Decision 8.", "")]),
    ])

    heading(doc, "Section 4 — Coping", 2)
    seed(doc, [
        ("Spine claim", [("Coping is a matching problem. What you can change determines what "
                          "you should try — and other people change what the trying costs.", "")]),
        ("Lead with", [("The two-step case. Regulate enough to act, then act. It is the "
                        "chapter's central skill and the lab already practises it.", "")]),
        ("Arc", [("You assume there is a right way to cope → it depends on what is "
                  "changeable, and on what state you are in → payoff.", "")]),
        ("Payoff line", [("There is no best coping strategy, the way there is no best tool. "
                          "There is only what the situation will let you change — and whether "
                          "you are currently in a state to change it.", "i")]),
        ("Social buffering seam", [("Chapter 10 said it already: for a while, the caregiver's "
                                    "regulated body is the infant's regulation system, borrowed "
                                    "from the outside. You never entirely stop borrowing. "
                                    "Reliable people do not only make a hard thing feel better "
                                    "— they change what the hard thing costs you to handle.", "i"),
                                   ("  (Decision 3)", "b")]),
        ("Prune", [("P4 — L121 entire, 27 words; the payoff line replaces its job. P5 — L177's "
                    "closing sentence, 17 words; the Hostinar review supersedes it. P3 — L76's "
                    "opening hedge and closing sentence, 22 words.", "")]),
        ("Watch", [("Do not use Hostinar to cash out the body budget. “Here is the "
                    "meter” is the overclaim the ledger bars, and it is the tempting "
                    "payoff for L55 and L76 too. Hostinar is a review — do not upgrade it to a "
                    "single demonstrated mechanism, and do not let “changes what it "
                    "costs” harden into a measured quantity. If the stranger contrast "
                    "comes up: Coan's own discussion says stranger contact probably does "
                    "regulate threat response in some people some of the time, so a flat "
                    "“strangers don't help” would be wrong, and unkind to a student "
                    "thinking about a counsellor.", "")]),
    ])

    heading(doc, "AI Connection", 2)
    seed(doc, [
        ("Spine claim", [("It produces the label without the body that made the label mean "
                          "anything.", "")]),
        ("Payoff line", [("It can say “overwhelmed” perfectly. It has nothing to be "
                          "overwhelmed about.", "i"),
                         ("  Then keep the existing close unchanged: ", ""),
                         ("Useful as a mirror; not proof of a felt state.", "i")]),
        ("Prune", [("The middle disclaimer sentence at L241. Recorded, not asked — "
                    "throat-clearing, and the neighbouring sentences carry the bound.", "")]),
        ("Watch", [("The ledger protects emotion-shaped output and the mirror line verbatim. "
                    "This lands the same no-stakes framing you chose for Chapter 11's AI "
                    "Connection, which is a consistency win rather than a coincidence.", "")]),
    ])

    heading(doc, "Closing cadence", 2)
    seed(doc, [
        ("Reclaim 4", [("An old, mostly ", ""), ("effective", "b"),
                        (" system returns to an old, mostly ", ""), ("excellent", "b"),
                        (" system. Dropping “prediction” was correct per the audit; "
                         "“excellent” was collateral.", "")]),
        ("Watch", [("The ledger protects the humane landing, the symbolic and chronic threat "
                    "landscape, and the short final rhythm. All three survive.", "")]),
    ])

    # ---------------- thread decisions ----------------
    heading(doc, "The three thread decisions", 1)
    decision(doc, 3, "Social buffering — install the Chapter 10 callback?", [
        "The seam is drafted in the Section 4 seed above, anchored on Hostinar, Sullivan and "
        "Gunnar (2014). See Findings 1 and 2 for why this anchor and not the fMRI study.",
        "Install as drafted / install with Thorsteinsson and James (1999) for an effect size "
        "instead / do not install, and delete the two Chapter 10 bridge claims so the chapter "
        "stops promising it.",
        [("If you decline it, say so and I will record what is lost: social support keeps its "
          "list of pathways and never acquires a mechanism, which was the spine's stated reason "
          "for routing the thread here.", "i")],
    ])
    decision(doc, 4, "Compression — install the core-affect-as-summary seam?", [
        "Drafted in the Section 1 seed. Compression is the spine's central lens and its "
        "drafting rules ask for it in every chapter; it currently scores zero here. Core "
        "affect's two numbers standing in for a whole body is the cleanest instance of lossy "
        "compression in the book, and it sets up granularity in the next section.",
        "Install / leave compression out of this chapter.",
    ])
    decision(doc, 5, "Action-readiness — name it or leave it implicit?", [
        "The chapter map assigns this chapter emotion as mattering and action-readiness. "
        "Action tendency currently appears in one table cell and one review answer, and nowhere "
        "in body prose.",
        "The Section 2 payoff line carries the idea without the term. Option A: let it ride "
        "unnamed. Option B: add the fear-anger-shame sentence and name it.",
    ])

    # ---------------- prunes ----------------
    heading(doc, "The prunes", 1)
    para(doc, "Every prune is yours. Total 184 words. Tick, strike, or amend.")
    datatable(doc, ["ID", "Where", "What comes out", "Words"],
              [["P1", "L88", "The constructionism meta-paragraph, whole", "59"],
               ["P2", "L55", "The two interchangeability examples", "22"],
               ["P3", "L76", "Opening hedge and closing sentence", "22"],
               ["P4", "L121", "Journaling, friendship, therapy — whole", "27"],
               ["P5", "L177", "The Heinrichs closing disclaimer", "17"],
               ["P6", "L137", "The goal-horizon scope clause", "20"],
               ["P7", "L69", "Closing sentence; restates the one before it", "17"]],
              widths=[0.45, 0.7, 4.6, 0.75])
    decision(doc, 7, "The seven prunes — and where the constructionism boundary lands", [
        "Approve, strike, or amend each of P1 through P7.",
        [("P1 has a consequence and needs a second answer.", "b"),
         (" It is the paragraph that states the constructionism boundary, and the audit "
          "requires that boundary be stated once. If P1 goes, it has to land somewhere. Three "
          "options: fold it into the opener's existing constructed-does-not-mean-chosen box; "
          "fold one clause into the theory table's constructionist row, which already says "
          "influential but debated; or keep P1 and cut it to a single sentence instead.", "")],
    ], answer_lines=4)

    # ---------------- reclaims ----------------
    heading(doc, "Reclaims from the repair pass", 1)
    para(doc, "Found by diffing the pre-repair and post-repair chapter. About 70 words. Most of "
              "the 4,900 words removed came out correctly; these four should not have.")
    datatable(doc, ["#", "What", "Status now", "Note"],
              [["1", "Right now, reading this, you are somewhere on that map — perhaps mildly "
                     "pleasant and moderately aroused; perhaps slightly unpleasant and "
                     "low-arousal.", "gone",
                "Makes core affect experiential instead of definitional. The raw-material "
                "sentence that followed it stays cut — that is the framework-specific claim."],
               ["2", "Kicks on when the house gets cold / learns your schedule and preheats the "
                     "house before you wake up", "flattened",
                "Must-preserve asset. Learns is also more accurate than knows — allostasis is "
                "acquired. Net zero words."],
               ["3", "A zebra's stress response mobilizes fully for the minute it takes to "
                     "escape a lion, then shuts off completely once the threat is gone or the "
                     "zebra is dead", "gone",
                "Voice-brief signature example, now only in Further Reading. Needs one word "
                "changed; see the Section 3 watch."],
               ["4", "An old, mostly excellent prediction system", "excellent lost",
                "Dropping prediction was right. Excellent was collateral."]],
              widths=[0.35, 2.5, 0.85, 2.8])
    para(doc, "These are register restorations, so they are yours — but they are also small and "
              "reversible, so they are one line rather than four boxes. Strike any you do not "
              "want.", size=10, color=GREY)

    # ---------------- must-preserve conflict ----------------
    heading(doc, "Must-preserve conflict", 1)
    decision(doc, 8, "The cognitive light cone's course-specific label", [
        "The audit's ledger requires the term keep its course-specific label. The repair pass "
        "replaced that with an attribution to Levin (2019) — which is correct, because Levin "
        "defines the term. So the ledger row and the spine's own description of it as a course "
        "metaphor rather than a standard research term are both overtaken.",
        "I am flagging rather than routing around it, as the protocol requires. Two answers "
        "needed: confirm the Levin attribution supersedes the ledger row, and say whether the "
        "spine's section 3 note gets amended to match. I have not touched the spine — that note "
        "also prescribes grounding citations, so it is a framework change and yours.",
    ], answer_lines=4)

    # ---------------- trade ----------------
    heading(doc, "The trade, measured", 1)
    para(doc, "Re-run 2026-07-26 against the current source, not quoted from the audit.",
         size=10, color=GREY)
    datatable(doc, ["Section", "Body words"],
              [["Opener", "283"], ["Section 1", "574"], ["Section 2", "839"],
               ["Section 3", "1,007"], ["Section 4", "616"], ["AI Connection", "130"],
               ["*Body total, opener to AI, excluding figure text", "*3,653"]],
              widths=[4.7, 1.8])
    datatable(doc, ["Addition", "Words (est.)"],
              [["Chapter 10 callback and Hostinar", "95"],
               ["Compression seam", "35"],
               ["Action-readiness seam", "30"],
               ["Four reclaims", "70"],
               ["Five payoff lines", "90"],
               ["*Total added", "*320"],
               ["*Total pruned", "*184"],
               ["*Net", "*+136"]],
              widths=[4.7, 1.8])
    para(doc, "Body 3,653 becomes 3,789, against a 3,500 to 5,000 target.")
    decision(doc, 9, "Net +136 words against a protocol rule that says cut instead", [
        [("The protocol says that if additions exceed prunes the seed set is too big and "
          "threads should be cut before presenting. I am presenting it anyway, and the "
          "reasoning should be checked rather than taken.", "b")],
        "That rule exists to stop an already-oversized chapter growing — Chapter 11 was 8,792 "
        "words when it was written. Chapter 12 is at 3,653 against its own target, so the "
        "rule's purpose does not apply, and satisfying its letter would mean cutting threads to "
        "shrink a chapter that is already short. Chapter 11 finished 22 per cent up with your "
        "agreement, on the grounds that what came out was catalog and what went in was "
        "narrative. The same argument applies here more easily.",
        "If you want strict net-negative anyway, the lever is dropping the compression seam and "
        "taking P1 and P4 wider. What is lost: compression is the one thing the spine's "
        "drafting rules ask of every chapter, and core affect's two numbers are its cleanest "
        "instance in the book.",
    ])

    # ---------------- integration ----------------
    heading(doc, "Integration, cheapest first", 1)
    bullet(doc, [("Invert lead-and-definition — free, highest yield. ", "b"),
                 ("Four sections already contain the right lead image arriving after the "
                  "definition it should precede: the thermostat, the granularity examples "
                  "(currently at the section's end), the snakes, and the zebra (currently "
                  "absent). The linter does not enforce term position; checked, not assumed.", "")])
    bullet(doc, [("Seam sentences — three, not ten. ", "b"),
                 ("Compression, action-readiness, social buffering. This chapter needs fewer "
                  "than Chapter 11 because its threads are already present.", "")])
    bullet(doc, [("Payoff lines — five, replacing pruned disclaimer positions. ", "b"),
                 ("This is the register work. It is not a separate pass.", "")])
    bullet(doc, [("Apparatus retrofit — free, and last. ", "b"),
                 ("One Stop and Retrieve across four sections, one Think About It, seven review "
                  "questions against a target of 8 to 12. Deferred by design so prompts test "
                  "the chapter as it finally reads. The audit's line stands: do not add another "
                  "question whose correct answer is constructed emotion.", "")])
    rich(doc, [("Sequencing, or the chapter gets rewritten twice: ", "b"),
               ("prune decisions, then lead and definition inversions, then reclaims, then "
                "seams, then payoff lines, then apparatus.", "")], before=6)
    rich(doc, [("The failure mode to name every time: ", "b"),
               ("adding a spine framing paragraph at the top of each section. It is the obvious "
                "move and it reproduces the spine-disconnected draft.", "")])

    # ---------------- recorded not asked ----------------
    heading(doc, "Recorded, not asked", 1)
    para(doc, "Decided without you, listed so you can overrule any of it.", size=10, color=GREY)
    for item in [
        "Trimming the middle disclaimer sentence in the AI Connection — throat-clearing, and "
        "the neighbouring sentences carry the bound.",
        "Dropping one sentence from the goal-horizon ladder.",
        "Moving the granularity examples from the end of Section 2 to its opening.",
        "Moving the snakes ahead of the lesion description in the Patient S.M. walkthrough.",
        "Not citing Grupe and Nitschke.",
        "Not restoring three pre-repair passages — this is why everything feels harder, the "
        "post-finals narrative, and social support predicts mortality. All three overclaimed. "
        "Their jobs are reassigned to payoff lines and to the Hostinar seam; the text stays cut.",
        "Not restoring the pre-repair body-budget review question. Correctly cut — it made the "
        "metaphor quasi-diagnostic.",
        "Not amending the spine's hedonic-treadmill note in this pass. It changes nothing here "
        "and belongs in a spine-maintenance pass.",
    ]:
        bullet(doc, [(item, "")])

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
