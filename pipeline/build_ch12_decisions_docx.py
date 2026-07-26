#!/usr/bin/env python3
"""Build the Chapter 12 settled-decisions record as DOCX.

This is a RECORD, not a review packet: all nine decisions were settled in
conversation on 2026-07-26. Nothing here has an answer field.

New filename by rule. The first review packet was overwritten in place by a
rebuild after the instructor had entered answers, and those answers were lost.
Any revision to a packet that has been sent goes to a new file.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "line-edit-packets" / "Chapter_12_Spine_Seeds_DECISIONS_2026-07-26.docx"

INK = RGBColor(0x11, 0x11, 0x11)
NAVY = RGBColor(0x0B, 0x25, 0x45)
GREY = RGBColor(0x55, 0x55, 0x55)
SETTLED = "E8F0E6"
REVISED = "FFF3D6"
NOTE = "E8EEF4"
SEED = "F7F7F5"


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), fill)
    tcPr.append(el)


def borders(table, sz=6, color="BBBBBB"):
    el = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color)
        el.append(e)
    table._tbl.tblPr.append(el)


def fix_widths(table, widths):
    table.autofit = False
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        table._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for wd in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(wd * 1440)))
        grid.append(gc)
    table._tbl.insert(1, grid)
    for row in table.rows:
        for i, wd in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(wd)


def run(p, text, *, bold=False, italic=False, size=10.5, color=INK):
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return r


def para(doc, text="", *, bold=False, italic=False, size=10.5, color=INK,
         after=6, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def rich(doc, segs, *, size=10.5, after=6, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    for t, s in segs:
        run(p, t, bold="b" in s, italic="i" in s, size=size)
    return p


def heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = NAVY
        r.font.name = "Calibri"
        r.font.size = Pt(16 if level == 1 else 12.5)
    h.paragraph_format.space_before = Pt(13)
    h.paragraph_format.space_after = Pt(4)
    return h


def box(doc, label, lines, fill):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(t)
    fix_widths(t, [6.5])
    c = t.cell(0, 0)
    shade(c, fill)
    first = c.paragraphs[0]
    first.paragraph_format.space_after = Pt(3)
    run(first, label, bold=True, size=10.5, color=NAVY)
    for line in lines:
        p = c.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        if isinstance(line, str):
            run(p, line, size=10)
        else:
            for t_, s_ in line:
                run(p, t_, bold="b" in s_, italic="i" in s_, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def dtable(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(t)
    for i, h in enumerate(headers):
        c = t.cell(0, i)
        shade(c, "E4E9EF")
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run(p, h, bold=True, size=9.5, color=NAVY)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            em = str(v).startswith("*")
            run(p, str(v)[1:] if em else str(v), bold=em, size=9.5)
    fix_widths(t, widths)
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(0.9)
        fp = s.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(fp, "Chapter 12 spine seeds — decisions settled 2026-07-26 — page ",
            size=8.5, color=GREY)
        r = fp.add_run()
        r.font.size = Pt(8.5)
        r.font.color.rgb = GREY
        b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
        it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " PAGE "
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
        for el in (b, it, e):
            r._r.append(el)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run(p, "Chapter 12 — Emotion, Stress & Coping", bold=True, size=21, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run(p, "Spine seeds — all nine decisions settled", size=13, color=GREY)

    box(doc, "What this is", [
        "A record, not a review packet. All nine decisions were settled in conversation on "
        "2026-07-26, so nothing here has an answer field. The diffable version is "
        "pipeline/ch12-spine-seeds.md, which also holds the reasoning for everything cut.",
        [("Why the filename changed. ", "b"),
         ("The first review packet was overwritten in place by a rebuild after you had entered "
          "answers, and those answers were lost. Any revision to a packet that has been sent "
          "now goes to a new file — the convention the Chapter 11 marked copy already followed.", "")],
        [("Before running git: delete C:\\GitHub\\psych101\\.git\\index.lock. ", "b"),
         ("A git status call stranded it; the sandbox mount cannot unlink files under .git.", "")],
    ], NOTE)

    heading(doc, "The nine decisions", 1)
    dtable(doc, ["#", "Decision", "Outcome"], [
        ["1", "Spine question", "Spine's wording plus a second clause. Replaces the "
         "table-of-contents sentence at the end of the opener."],
        ["2", "Chapter engine", "Accepted as drafted. Drafting test, not chapter text."],
        ["3", "Social buffering", "*Installed. Hostinar, Sullivan & Gunnar (2014), not the "
         "2006 fMRI study."],
        ["4", "Compression", "*Thread declined; sentence kept. Framework gesture cut."],
        ["5", "Action-readiness", "Payoff line only; term not named."],
        ["6", "Patient S.M.", "Installed, first version. Residual risk accepted."],
        ["7", "The prunes", "*Settled — three of seven revised against my own "
         "recommendation."],
        ["8", "Light cone label", "Levin supersedes the audit. Spine section 3 amended."],
        ["9", "Net words", "Moot. Net is +117 on a chapter at the low end of target."],
    ], [0.35, 1.5, 4.65])

    heading(doc, "The four that changed during discussion", 1)
    para(doc, "These are the ones where the packet's recommendation did not survive "
              "scrutiny. Recorded because the next pass should not re-propose them.",
         size=10, color=GREY)

    box(doc, "Decision 4 — compression thread declined", [
        [("Your question — ", ""), ("“4 only to get compression in?”", "i"),
         (" — was correct, and the answer was yes. That is the same reason just-world and "
          "attachment were cut from Chapter 11: the payoff was framework coherence, not "
          "student comprehension.", "")],
        [("Kept, on student payoff alone: ", "b"),
         ("Two numbers — pleasant or not, wound up or not — standing in for everything "
          "happening in a body. It is why you can be certain you feel terrible and have no "
          "idea why.", "i")],
        [("Cut as framework tidiness: ", "b"),
         ("That is a summary, and summaries leave things out.", "i"),
         (" Meta-commentary — the same species as the throat-clearing being removed "
          "elsewhere.", "")],
        "A second reason, missed on the first pass: the spine's own drafting rules forbid "
        "collapsing prediction error, emotion, and stress into one mechanism. Calling core "
        "affect compression implies emotion construction is the same operation as "
        "episodic-to-semantic memory compression, and it is not.",
        [("What is lost: ", "b"),
         ("Chapter 12 stays the only chapter with no compression thread. Assessed as correct "
          "rather than as a gap. Flagged as a spine-maintenance question — the instruction to "
          "use the lens in every chapter may be overreaching.", "")],
    ], REVISED)

    box(doc, "Decision 7 — three prunes revised, two deferred then unblocked", [
        [("P2 ", "b"), ("— the packet proposed cutting the two interchangeability examples. "
          "Reversed: the paragraph states its bound four times and those examples are the only "
          "concrete instance of “not one currency” in the chapter. The closing "
          "sentence goes instead. Same saving.", "")],
        [("P1 ", "b"), ("— the packet proposed folding the constructionism boundary into the "
          "theory table. Rejected: that buries the chapter's central epistemic commitment in a "
          "table cell students skim. Instead the paragraph is cut to that one sentence, in "
          "prose, at the head of Section 2.", "")],
        [("P3 ", "b"), ("— the anti-diagnosis clause is relocated to the existing Think About "
          "It prompt, not deleted. Same protection, out of the payoff position. Net saving is "
          "the 8-word opening hedge only.", "")],
        [("P4 and P5 ", "b"), ("were listed as independent prunes. They are the second half of "
          "Decision 3 — cutting P5 without installing social buffering would leave Heinrichs as "
          "an unbounded general claim, which is worse than the current disclaimer. Both "
          "unblocked now that Decision 3 is in.", "")],
        [("P6 ", "b"), ("depended on the zebra. With the zebra leading Section 3, the three "
          "goal scales are no longer the paragraph's structure, so the anti-hierarchy warning "
          "is no longer protecting against a live misreading.", "")],
        [("Settled total is 149 words, not the 184 the packet claimed", "b"),
         (" — that figure assumed the maximalist P1 and P2.", "")],
    ], REVISED)

    box(doc, "Decision 5 — cheaper than budgeted", [
        "The granularity bullets already are action tendencies, and “That precision "
        "can suggest different actions” is already the payoff, flat and hedged. So this "
        "is a replacement, not an addition: about 31 words for 6, net +25 rather than the +30 "
        "budgeted. The optional naming sentence was not taken — the chapter map wants "
        "action-readiness taught, not the term recited.",
    ], REVISED)

    box(doc, "Decision 8 — an audit requirement retired, and the spine corrected", [
        "The audit required the cognitive light cone keep its course-specific label. Levin "
        "(2019) defines the term — the outer boundary in space and time of the largest goal a "
        "system can work toward — so the chapter's attribution is more accurate than the "
        "document protecting it.",
        [("theoretical-spine.md section 3 amended at your direction. ", "b"),
         ("It had called the phrase a course-specific metaphor and not a standard research "
          "term. It now attributes it to Levin, states the correction, keeps Suddendorf, "
          "Schacter, Gilbert & Wilson and Sapolsky as grounding for the application, and "
          "records that Levin derives the expanding boundary from a homeostatic drive to reduce "
          "stress — which links the light cone back to Section 1's allostasis material instead "
          "of leaving it stranded in Section 3.", "")],
        "Grupe & Nitschke closed in the same edit: the citation verifies, but it is about "
        "pathological anxiety, so it imports a diagnostic frame the audit bars. Declined for "
        "Chapter 12, rerouted as a Chapter 13 candidate.",
    ], REVISED)

    heading(doc, "What goes into the chapter", 1)
    para(doc, "Sequenced as the protocol requires: prunes, then lead-and-definition "
              "inversions, then reclaims, then seams, then payoff lines, then apparatus last.",
         size=10, color=GREY)

    for title, lines in [
        ("Spine question — end of the opener", [
            [("So how does your body decide what matters — and why won't it always stand "
              "down?", "i")],
            "The preceding sentence already contains “what matters,” so the "
            "question picks up a word the paragraph just used. Replaces "
            "“This chapter follows that sequence from regulation to emotion to stress "
            "to action.”",
            "This makes Chapter 12 the fifth of thirteen chapters to carry its spine question "
            "in the text. The other eight are a backlog item, not a Chapter 12 one.",
        ]),
        ("Opener — payoff line", [
            [("The alarm in your kitchen reacts to smoke. You react to what smoke might mean "
              "— and nobody arrives at that with the same history.", "i")],
        ]),
        ("Section 1 — lead, payoff, and two reclaims", [
            "Lead with the thermostat, verbs restored: kicks on when the house gets cold, and "
            "learns your schedule and preheats the house before you wake up. Learns is also "
            "more accurate than knows, because allostasis is acquired.",
            [("By the time you feel anything, your body has already moved. Feeling is not the "
              "report on what happened — it is part of getting ready for what might.", "i")],
            [("Right now, reading this, you are somewhere on that map — perhaps mildly pleasant "
              "and moderately aroused; perhaps slightly unpleasant and low-arousal.", "i")],
            [("Two numbers — pleasant or not, wound up or not — standing in for everything "
              "happening in a body. It is why you can be certain you feel terrible and have no "
              "idea why.", "i")],
        ]),
        ("Section 2 — lead and payoff", [
            "Move the granularity examples from the end of the section to its opening. They are "
            "the concrete hook and they currently arrive after the theory table and S.M.",
            [("“I feel bad” tells you nothing about what to do next. "
              "“I'm disappointed” tells you to revise an expectation. "
              "“I'm frustrated” tells you to change the approach. Same discomfort, "
              "different instructions.", "i")],
            "Section 2 keeps one sentence of its opening paragraph, in prose, at the head: "
            "constructionism is the lens used here, but it is not the only serious contemporary "
            "account.",
        ]),
        ("Patient S.M. — lead and payoff", [
            "Move the snakes ahead of the lesion description.",
            [("She can tell you the snake is dangerous. She just doesn't mind. Knowing and "
              "minding turn out to be two different jobs — and the CO2 result shows the "
              "amygdala is not the only way into the second one. The amygdala contributes "
              "importantly to some externally triggered threats; it is not necessary for every "
              "experience of fear or panic.", "i")],
            [("Residual risk, accepted knowingly. ", "b"),
             ("“She just doesn't mind” is true of the snakes, films and haunted "
              "house, and would be false as a general statement about her — she panicked on "
              "CO2. The following sentence has to carry that limit. If it reads as a general "
              "claim at line-edit, the conservative alternative is in the markdown record.", "")],
        ]),
        ("Section 3 — reorder to zebra, exam, definitions, light cone", [
            "The section currently opens on four bolded terms in four sentences, with the exam "
            "image second. The definitions should arrive to name what the student is already "
            "looking at.",
            [("The zebra's stress response mobilizes fully for the minute it takes to escape a "
              "lion, then shuts off completely once the threat is gone or the zebra is dead "
              "(Sapolsky, 2004). Yours doesn't. Your threats are next Tuesday — and they will "
              "still be next Tuesday tomorrow.", "i")],
            [("One word changed from the pre-repair version. ", "b"),
             ("It said humans run the same ancient mobilization machinery. The audit bars "
              "identical-pathway language and that is the same claim. It becomes overlapping, "
              "which the chapter already says correctly elsewhere.", "")],
            "The spine already prescribed Sapolsky's zebra contrast as grounding for this "
            "material. The repair pass removed the only place the chapter used it, so this "
            "restores something the framework had asked for.",
        ]),
        ("Section 4 — payoff and the social buffering seam", [
            [("There is no best coping strategy, the way there is no best tool. There is only "
              "what the situation will let you change — and whether you are currently in a "
              "state to change it.", "i")],
            [("Chapter 10 said it already: for a while, the caregiver's regulated body is the "
              "infant's regulation system, borrowed from the outside. You never entirely stop "
              "borrowing. Reliable people do not only make a hard thing feel better — they "
              "change what the hard thing costs you to handle (Hostinar, Sullivan, & Gunnar, "
              "2014).", "i")],
            "The Chapter 10 clause is quoted verbatim from that chapter's own sentence. It "
            "replaces “Those are several pathways, not one deposit into a literal "
            "account,” keeping Cohen & Wills and the emotion- and problem-focused closer "
            "either side of it.",
            [("Watch: ", "b"),
             ("“changes what the hard thing costs” must stay a claim about "
              "regulatory demand, never about a measurable balance. Hostinar is a review — do "
              "not upgrade it to a single demonstrated mechanism.", "")],
        ]),
        ("AI Connection and the closing cadence", [
            [("It can say “overwhelmed” perfectly. It has nothing to be "
              "overwhelmed about.", "i"),
             ("  Then the existing close, unchanged: ", ""),
             ("Useful as a mirror; not proof of a felt state.", "i")],
            "Closing cadence: an old, mostly effective system returns to an old, mostly "
            "excellent system. Dropping “prediction” was correct per the audit; "
            "“excellent” was collateral.",
        ]),
    ]:
        box(doc, title, lines, SEED)

    heading(doc, "The trade, measured", 1)
    dtable(doc, ["Insertion", "Gross", "Replaces", "Net"], [
        ["Spine question", "16", "14", "+2"],
        ["Opener payoff line", "25", "—", "+25"],
        ["Section 1 payoff line", "29", "—", "+29"],
        ["Reclaim — core-affect second person", "22", "—", "+22"],
        ["Decision 4 — kept sentence", "34", "—", "+34"],
        ["Section 2 payoff / Decision 5", "31", "6", "+25"],
        ["Decision 6 — S.M. block", "59", "62", "−3"],
        ["Reclaim — zebra lead", "47", "—", "+47"],
        ["Section 4 payoff line", "35", "—", "+35"],
        ["Decision 3 — social buffering", "53", "15", "+38"],
        ["AI Connection payoff line", "12", "—", "+12"],
        ["*Total added", "", "", "*+266"],
        ["*Total pruned", "", "", "*−149"],
        ["*Net", "", "", "*+117"],
    ], [3.4, 0.85, 1.0, 1.25])
    para(doc, "Body 3,653 becomes 3,770, against a 3,500 to 5,000 target. Prunes are exact "
              "strings; additions will move during drafting, so treat these as the budget "
              "rather than the outcome.")
    rich(doc, [("Correction on the record. ", "bi"),
               ("An intermediate figure of +46 was stated in conversation. That was a guess "
                "made in prose, not a calculation. Third measurement correction in this pass — "
                "the pattern is that numbers stated without running them are wrong, and the "
                "instruction to re-run rather than remember applies to the agent's own "
                "arithmetic, not only to the audit's.", "i")])

    heading(doc, "Open, and not part of this pass", 1)
    for t in [
        "Ecological dominance and social competition was added to the spine trunk by a parallel "
        "session the same day, and the backlog says Chapter 12 gains the ultimate ground for "
        "social-evaluative stress. Your own paper — Oxford, Ponzi & Geary (2010) — is a "
        "cortisol finding about social-evaluative threat, and this chapter invokes socially "
        "evaluated threat three times without saying why social evaluation is threatening. Left "
        "out because the backlog sequences proximate-and-ultimate into Chapter 1 first. Your "
        "call whether Chapter 12 waits.",
        "Voice loss is systematic, not a Chapter 12 problem: throat-clearing rose in 11 of 11 "
        "audited chapters and not one fell. Chapter 13 is next in the queue and third worst, "
        "with paragraphs landing on a negation doubled from 22 to 44 per cent. Chapter 5 and "
        "Chapter 3 outrank several chapters ahead of them.",
        "HANDOFF.md still lists settling the packet format as a next action. "
        "spine-seed-protocol.md still calls it unresolved, and its delivery format still does "
        "not say that a sent packet is never rebuilt in place, or that decision boxes need the "
        "current text quoted beside the proposal.",
        "CLAUDE.md says read-only git commands are safe under the sandbox mount. That is wrong "
        "for git status, which strands an index lock.",
    ]:
        b = doc.add_paragraph(style="List Bullet")
        b.paragraph_format.space_after = Pt(4)
        b.paragraph_format.left_indent = Inches(0.4)
        run(b, t, size=10)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
