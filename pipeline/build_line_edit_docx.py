#!/usr/bin/env python3
"""Build lightweight Word line-edit packets from canonical chapter Markdown.

The Markdown remains authoritative. These DOCX files are disposable editorial
snapshots intended for Track Changes and comments, not publication artifacts.

Design preset: compact_reference_guide.
Named overrides:
  * line_edit_title: Calibri 24 pt, #0B2545, bold, 0/4 pt spacing.
  * editorial_note: Calibri 9 pt, #555555, 6/12 pt spacing, light-gray shade.
  * artwork_placeholder: Calibri 9 pt italic, #666666, 3/3 pt spacing.
"""

from __future__ import annotations

import argparse
import html
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "source" / "chapters"
DEFAULT_OUTPUT = ROOT / "line-edit-packets"

CHAPTER_SOURCES = {
    3: "ch03-neuroscience-biological-bases.md",
    4: "ch04-sensation-perception.md",
    5: "ch05-consciousness.md",
    6: "ch06-sleep.md",
    7: "ch07-learning.md",
    8: "ch08-memory.md",
    9: "09-thinking-language-intelligence.md",
    10: "ch10-lifespan-development.md",
    11: "ch11-social-psychology.md",
    12: "ch12-emotion-stress-coping.md",
    13: "ch13-psychological-disorders-therapy.md",
}

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x0B, 0x25, 0x45)
GRAY = RGBColor(0x55, 0x55, 0x55)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        element = tc_mar.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_fixed_table_geometry(table, widths):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (fld_char, instr, separate, text, end):
        run._r.append(node)


def set_font(run, size=11, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_numbering_definition(doc, kind):
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing_abstract, default=-1) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for level in range(9):
        left = 540 + level * 360
        hanging = 270
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else f"%{level + 1}.")
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, suff, lvl_jc, p_pr])
        abstract.append(lvl)

    numbering.append(abstract)
    nums = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(nums, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_paragraph_numbering(paragraph, num_id, level=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 8)))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def clean_inline_markup(text):
    text = html.unescape(text)
    text = re.sub(r"<span\b[^>]*>(.*?)</span>", r"**\1**", text, flags=re.I)
    text = re.sub(r"<br\s*/?>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
    text = text.replace("\\|", "|")
    return text.strip()


INLINE_PATTERN = re.compile(
    r"(\*\*\*.+?\*\*\*|___.*?___|\*\*.+?\*\*|__.*?__|(?<!\*)\*[^*\n]+?\*(?!\*)|(?<!_)_[^_\n]+?_(?!_)|`[^`]+`)"
)


def add_inline(paragraph, text, *, default_bold=False, default_italic=False, color=None, size=11):
    text = clean_inline_markup(text)
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_font(run, size=size, color=color, bold=default_bold, italic=default_italic)
        token = match.group(0)
        bold = default_bold
        italic = default_italic
        mono = False
        if token.startswith(("***", "___")):
            content = token[3:-3]
            bold = italic = True
        elif token.startswith(("**", "__")):
            content = token[2:-2]
            bold = True
        elif token.startswith(("*", "_")):
            content = token[1:-1]
            italic = True
        else:
            content = token[1:-1]
            mono = True
        run = paragraph.add_run(content)
        set_font(run, size=size, color=color, bold=bold, italic=italic, name="Consolas" if mono else "Calibri")
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font(run, size=size, color=color, bold=default_bold, italic=default_italic)


def shade_paragraph(paragraph, fill="F4F6F9", left_border="2E74B5"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "14")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), left_border)
    borders.append(left)
    p_pr.append(borders)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "180")
    ind.set(qn("w:right"), "120")
    p_pr.append(ind)


def split_table_row(line):
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in re.split(r"(?<!\\)\|", stripped)]


def is_table_separator(cells):
    return cells and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def add_markdown_table(doc, rows):
    parsed = [split_table_row(row) for row in rows]
    if len(parsed) > 1 and is_table_separator(parsed[1]):
        parsed.pop(1)
    if not parsed:
        return
    col_count = max(len(row) for row in parsed)
    parsed = [row + [""] * (col_count - len(row)) for row in parsed]

    table = doc.add_table(rows=len(parsed), cols=col_count)
    table.style = "Table Grid"
    widths = [9360 // col_count] * col_count
    widths[-1] += 9360 - sum(widths)
    set_fixed_table_geometry(table, widths)
    for row_idx, row in enumerate(parsed):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            if row_idx == 0:
                set_cell_shading(cell, "E8EEF5")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.15
            add_inline(paragraph, value, default_bold=(row_idx == 0), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def preprocess_markdown(text):
    text = re.sub(r"<!--.*?-->", "", text, flags=re.S)
    text = re.sub(
        r"<video\b.*?</video>",
        "\n[Video artwork omitted from line-edit packet; caption follows.]\n",
        text,
        flags=re.S | re.I,
    )
    lines = text.replace("\r\n", "\n").splitlines()
    # Sources place non-student provenance/routing notes in blockquotes between
    # the H1 title and the first horizontal rule. They do not belong in a line-
    # edit copy, while later blockquotes are real student-facing callouts.
    before_first_rule = True
    filtered = []
    for line in lines:
        if before_first_rule and line.strip() == "---":
            before_first_rule = False
        if before_first_rule and line.lstrip().startswith(">"):
            continue
        filtered.append(line)
    return "\n".join(filtered)


def emit_paragraph(doc, text):
    text = " ".join(part.strip() for part in text if part.strip())
    if not text:
        return
    paragraph = doc.add_paragraph()
    add_inline(paragraph, text)


def markdown_to_docx(doc, markdown_text):
    lines = preprocess_markdown(markdown_text).splitlines()
    body = []
    bullet_num_id = None
    decimal_num_id = None
    active_list_kind = None
    in_code = False
    code_lines = []
    i = 0

    def flush_body():
        nonlocal body
        if body:
            emit_paragraph(doc, body)
            body = []

    def reset_list():
        nonlocal active_list_kind, bullet_num_id, decimal_num_id
        active_list_kind = None
        bullet_num_id = None
        decimal_num_id = None

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if stripped.startswith("```"):
            flush_body()
            reset_list()
            if in_code:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.space_after = Pt(6)
                run = paragraph.add_run("\n".join(code_lines))
                set_font(run, size=9, color=GRAY, name="Consolas")
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not stripped:
            flush_body()
            reset_list()
            i += 1
            continue

        if stripped == "---":
            flush_body()
            reset_list()
            i += 1
            continue

        if stripped.startswith("<") and stripped.endswith(">") and not re.search(r">[^<]+<", stripped):
            flush_body()
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_body()
            reset_list()
            level = len(heading.group(1))
            text = clean_inline_markup(heading.group(2))
            if level == 1:
                # The chapter title is already used in the packet title block.
                i += 1
                continue
            style = "Heading 1" if level == 2 else "Heading 2" if level == 3 else "Heading 3"
            paragraph = doc.add_paragraph(style=style)
            add_inline(paragraph, text, default_bold=True, color=BLUE if style != "Heading 3" else DARK_BLUE,
                       size=16 if style == "Heading 1" else 13 if style == "Heading 2" else 12)
            i += 1
            continue

        if re.match(r"^!\[[^\]]*\]\([^)]*\)\s*$", stripped):
            flush_body()
            reset_list()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run("[Figure artwork omitted from line-edit packet; caption follows.]")
            set_font(run, size=9, color=RGBColor(0x66, 0x66, 0x66), italic=True)
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_body()
            reset_list()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        if stripped.startswith(">"):
            flush_body()
            reset_list()
            quote_parts = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                part = re.sub(r"^\s*>\s?", "", lines[i]).strip()
                if part:
                    quote_parts.append(part)
                i += 1
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            shade_paragraph(paragraph)
            add_inline(paragraph, " ".join(quote_parts))
            continue

        list_match = re.match(r"^(\s*)([-+*]|\d+\.)\s+(.+)$", raw)
        if list_match:
            flush_body()
            indent_spaces = len(list_match.group(1).replace("\t", "    "))
            level = min(indent_spaces // 2, 8)
            marker = list_match.group(2)
            kind = "bullet" if marker in ("-", "+", "*") else "decimal"
            if active_list_kind != kind:
                active_list_kind = kind
                if kind == "bullet":
                    bullet_num_id = add_numbering_definition(doc, "bullet")
                else:
                    decimal_num_id = add_numbering_definition(doc, "decimal")
            paragraph = doc.add_paragraph(style="List Bullet" if kind == "bullet" else "List Number")
            set_paragraph_numbering(paragraph, bullet_num_id if kind == "bullet" else decimal_num_id, level)
            add_inline(paragraph, list_match.group(3))
            i += 1
            continue

        # Skip provenance-only editorial routing notes at the top of sources.
        if "Drafting history & provenance" in stripped or "Drafting history and provenance" in stripped:
            flush_body()
            i += 1
            continue

        if stripped.startswith("[") and "omitted from line-edit packet" in stripped:
            flush_body()
            reset_list()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(stripped)
            set_font(run, size=9, color=RGBColor(0x66, 0x66, 0x66), italic=True)
            i += 1
            continue

        body.append(stripped)
        i += 1

    flush_body()


def extract_title(markdown_text, chapter_number):
    match = re.search(r"^#\s+(.+)$", markdown_text, flags=re.M)
    if match:
        return clean_inline_markup(match.group(1))
    return f"Chapter {chapter_number}"


def safe_filename(title):
    name = re.sub(r"^Chapter\s+(\d+)\s*:\s*", r"Chapter_\1_", title, flags=re.I)
    name = re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_")
    return name + "_Line_Edit.docx"


def build_packet(chapter_number, source_path, output_dir):
    markdown_text = source_path.read_text(encoding="utf-8")
    title = extract_title(markdown_text, chapter_number)
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run(f"Chapter {chapter_number} line-edit copy")
    set_font(run, size=8.5, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Page ")
    set_font(run, size=8.5, color=GRAY)
    add_page_field(footer)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.paragraph_format.keep_with_next = True
    title_run = title_p.add_run(title)
    set_font(title_run, size=24, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)
    sub_run = subtitle.add_run("Instructor line-edit copy")
    set_font(sub_run, size=12, color=GRAY, italic=True)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(12)
    shade_paragraph(note, fill="F2F4F7", left_border="A6A6A6")
    add_inline(
        note,
        f"Snapshot generated {date.today().isoformat()} from {source_path.relative_to(ROOT)}. "
        "Use Track Changes or comments for editorial review. Figures and video artwork are omitted, "
        "but their captions remain. The Markdown source remains authoritative.",
        color=GRAY,
        size=9,
    )

    markdown_to_docx(doc, markdown_text)

    core = doc.core_properties
    core.title = title + " - Line Edit"
    core.subject = "Psychology 101 instructor line-edit packet"
    core.author = "Oxford Psychology 101"
    core.comments = "Disposable editorial snapshot generated from canonical Markdown."

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / safe_filename(title)
    doc.save(output_path)
    return output_path


def parse_args():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--chapters",
        nargs="+",
        type=int,
        default=sorted(CHAPTER_SOURCES),
        choices=sorted(CHAPTER_SOURCES),
        help="Chapter numbers to export (default: 3-13).",
    )
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT)
    return parser.parse_args()


def main():
    args = parse_args()
    for chapter in args.chapters:
        source_path = SOURCE_DIR / CHAPTER_SOURCES[chapter]
        output = build_packet(chapter, source_path, args.output_dir)
        print(output)


if __name__ == "__main__":
    main()
